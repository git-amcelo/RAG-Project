"""
Document Processing Module for RAG System
Weeks 4-5: PDF and Word document processing with intelligent chunking

Implements:
- PDF text extraction using PyPDF2 and pdfplumber
- Word document processing using python-docx
- Intelligent text chunking with LangChain
- Metadata extraction and preservation
"""

import os
import time
import hashlib
from typing import Dict, List, Optional, Tuple
from datetime import datetime
from pathlib import Path
import warnings
warnings.filterwarnings('ignore')

try:
    import PyPDF2
    import pdfplumber
except ImportError:
    print("Installing PDF processing libraries...")
    os.system("pip install PyPDF2 pdfplumber")

try:
    from docx import Document
except ImportError:
    print("Installing Word processing library...")
    os.system("pip install python-docx")

try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_core.documents import Document as LangDocument
except ImportError:
    print("Using alternative imports...")
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.schema import Document as LangDocument

from sentence_transformers import SentenceTransformer


class DocumentProcessor:
    """
    Process PDF and Word documents for the RAG system

    Features:
    - Multi-format support (PDF, DOCX)
    - Intelligent text chunking
    - Metadata extraction
    - Progress tracking
    """

    def __init__(self,
                 chunk_size: int = 1000,
                 chunk_overlap: int = 200,
                 chunk_batch_size: int = 32):
        """
        Initialize document processor

        Args:
            chunk_size: Maximum characters per chunk
            chunk_overlap: Character overlap between chunks
            chunk_batch_size: Batch size for processing
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.chunk_batch_size = chunk_batch_size

        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )

        # Supported formats
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx
        }

        # Storage paths
        self.upload_dir = Path("data/uploads")
        self.processed_dir = Path("data/documents")

        # Create directories
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self.processed_dir.mkdir(parents=True, exist_ok=True)

        print(f"✓ DocumentProcessor initialized")
        print(f"  Chunk size: {chunk_size}, Overlap: {chunk_overlap}")
        print(f"  Upload directory: {self.upload_dir}")
        print(f"  Processed directory: {self.processed_dir}")

    def validate_file(self, file_path: str) -> Tuple[bool, Optional[str]]:
        """
        Validate uploaded file

        Args:
            file_path: Path to uploaded file

        Returns:
            Tuple of (is_valid, error_message)
        """
        path = Path(file_path)

        # Check if file exists
        if not path.exists():
            return False, "File does not exist"

        # Check file extension
        if path.suffix.lower() not in self.supported_formats:
            return False, f"Unsupported format: {path.suffix}. Supported: {list(self.supported_formats.keys())}"

        # Check file size (max 50MB)
        file_size = path.stat().st_size
        if file_size > 50 * 1024 * 1024:
            return False, f"File too large: {file_size / (1024*1024):.1f}MB (max 50MB)"

        return True, None

    def process_document(self, file_path: str, document_id: Optional[str] = None) -> Dict:
        """
        Process a document end-to-end

        Args:
            file_path: Path to document file
            document_id: Optional document ID (auto-generated if not provided)

        Returns:
            Processing result with chunks and metadata
        """
        start_time = time.time()

        # Validate file
        is_valid, error_msg = self.validate_file(file_path)
        if not is_valid:
            return {
                "success": False,
                "error": error_msg,
                "document_id": document_id
            }

        # Generate document ID if not provided
        if document_id is None:
            document_id = self._generate_document_id(file_path)

        # Get file info
        path = Path(file_path)
        file_extension = path.suffix.lower()

        print(f"\nProcessing document: {path.name}")
        print(f"  Format: {file_extension}")
        print(f"  Document ID: {document_id}")

        try:
            # Extract text based on format
            process_func = self.supported_formats[file_extension]
            text_content, metadata = process_func(file_path, document_id)

            if not text_content or len(text_content.strip()) < 50:
                return {
                    "success": False,
                    "error": "Extracted text too short (min 50 characters)",
                    "document_id": document_id
                }

            # Split into chunks
            chunks = self._split_into_chunks(text_content, document_id, metadata)

            processing_time = time.time() - start_time

            result = {
                "success": True,
                "document_id": document_id,
                "file_name": path.name,
                "file_type": file_extension,
                "file_size": path.stat().st_size,
                "total_characters": len(text_content),
                "num_chunks": len(chunks),
                "processing_time": processing_time,
                "chunks": chunks,
                "metadata": metadata,
                "processed_at": datetime.now().isoformat()
            }

            print(f"✓ Document processed successfully")
            print(f"  Extracted {len(text_content)} characters")
            print(f"  Created {len(chunks)} chunks")
            print(f"  Processing time: {processing_time:.2f}s")

            return result

        except Exception as e:
            print(f"❌ Error processing document: {e}")
            import traceback
            traceback.print_exc()

            return {
                "success": False,
                "error": str(e),
                "document_id": document_id
            }

    def _process_pdf(self, file_path: str, document_id: str) -> Tuple[str, Dict]:
        """
        Extract text from PDF using both PyPDF2 and pdfplumber

        Args:
            file_path: Path to PDF file
            document_id: Document identifier

        Returns:
            Tuple of (extracted_text, metadata)
        """
        text_content = []
        metadata = {
            "document_id": document_id,
            "source_type": "pdf",
            "pages": [],
            "extraction_method": []
        }

        # Try PyPDF2 first
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)

                for page_num, page in enumerate(pdf_reader.pages, start=1):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(page_text)
                            metadata["pages"].append({
                                "page_number": page_num,
                                "character_count": len(page_text)
                            })
                    except Exception as e:
                        print(f"  Warning: Could not extract page {page_num}: {e}")

                metadata["extraction_method"].append("PyPDF2")
                metadata["total_pages"] = num_pages

        except Exception as e:
            print(f"  PyPDF2 extraction failed: {e}")

        # Try pdfplumber as backup/additional
        try:
            with pdfplumber.open(file_path) as pdf:
                if not metadata["extraction_method"]:
                    metadata["total_pages"] = len(pdf.pages)

                for page_num, page in enumerate(pdf.pages, start=1):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            if page_num <= len(text_content):
                                # Combine with existing text
                                text_content[page_num - 1] += " " + page_text
                            else:
                                text_content.append(page_text)
                    except Exception as e:
                        print(f"  Warning: pdfplumber could not extract page {page_num}: {e}")

                metadata["extraction_method"].append("pdfplumber")

        except Exception as e:
            print(f"  pdfplumber extraction failed: {e}")

        # Combine all pages
        full_text = "\n\n".join(text_content)

        return full_text, metadata

    def _process_docx(self, file_path: str, document_id: str) -> Tuple[str, Dict]:
        """
        Extract text from Word document

        Args:
            file_path: Path to DOCX file
            document_id: Document identifier

        Returns:
            Tuple of (extracted_text, metadata)
        """
        doc = Document(file_path)

        paragraphs = []
        metadata = {
            "document_id": document_id,
            "source_type": "docx",
            "paragraphs": [],
            "total_paragraphs": 0
        }

        for para_num, paragraph in enumerate(doc.paragraphs, start=1):
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)
                metadata["paragraphs"].append({
                    "paragraph_number": para_num,
                    "character_count": len(text),
                    "style": paragraph.style.name if paragraph.style else "Normal"
                })

        metadata["total_paragraphs"] = len(paragraphs)

        # Combine paragraphs
        full_text = "\n\n".join(paragraphs)

        return full_text, metadata

    def _split_into_chunks(self, text: str, document_id: str, metadata: Dict) -> List[Dict]:
        """
        Split text into intelligent chunks using LangChain

        Args:
            text: Full text to chunk
            document_id: Document identifier
            metadata: Document metadata

        Returns:
            List of chunk dictionaries
        """
        # Create LangChain document
        lang_doc = LangDocument(page_content=text, metadata=metadata.copy())

        # Split into chunks
        chunk_texts = self.text_splitter.split_documents([lang_doc])

        chunks = []
        for chunk_num, chunk_doc in enumerate(chunk_texts, start=1):
            chunk_text = chunk_doc.page_content
            chunk_metadata = chunk_doc.metadata.copy()

            # Add chunk-specific metadata
            chunk_metadata.update({
                "chunk_id": f"{document_id}_chunk_{chunk_num:04d}",
                "chunk_number": chunk_num,
                "character_count": len(chunk_text),
                "document_id": document_id
            })

            chunks.append({
                "chunk_id": chunk_metadata["chunk_id"],
                "chunk_number": chunk_num,
                "text": chunk_text,
                "metadata": chunk_metadata
            })

        return chunks

    def _generate_document_id(self, file_path: str) -> str:
        """
        Generate unique document ID based on file content

        Args:
            file_path: Path to file

        Returns:
            Unique document ID
        """
        path = Path(file_path)

        # Create hash from file name and timestamp
        hash_input = f"{path.name}_{time.time()}_{path.stat().st_size}"
        hash_value = hashlib.md5(hash_input.encode()).hexdigest()[:12]

        # Create readable ID
        safe_name = path.stem.replace(" ", "_").replace("-", "_")[:20]
        document_id = f"doc_{safe_name}_{hash_value}"

        return document_id

    def save_processed_document(self, result: Dict) -> str:
        """
        Save processed document to disk

        Args:
            result: Processing result dictionary

        Returns:
            Path to saved file
        """
        if not result.get("success"):
            return None

        document_id = result["document_id"]
        save_path = self.processed_dir / f"{document_id}.json"

        import json

        with open(save_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=2, ensure_ascii=False)

        print(f"✓ Saved processed document to {save_path}")
        return str(save_path)

    def get_processed_document(self, document_id: str) -> Optional[Dict]:
        """
        Load processed document from disk

        Args:
            document_id: Document identifier

        Returns:
            Processed document dictionary or None
        """
        load_path = self.processed_dir / f"{document_id}.json"

        if not load_path.exists():
            return None

        import json

        with open(load_path, 'r', encoding='utf-8') as f:
            result = json.load(f)

        return result

    def list_processed_documents(self) -> List[Dict]:
        """
        List all processed documents

        Returns:
            List of document metadata
        """
        documents = []

        for file_path in self.processed_dir.glob("*.json"):
            try:
                import json
                with open(file_path, 'r', encoding='utf-8') as f:
                    result = json.load(f)

                if result.get("success"):
                    documents.append({
                        "document_id": result["document_id"],
                        "file_name": result["file_name"],
                        "file_type": result["file_type"],
                        "num_chunks": result["num_chunks"],
                        "processed_at": result["processed_at"],
                        "file_size": result["file_size"]
                    })
            except Exception as e:
                print(f"Warning: Could not load {file_path}: {e}")

        return documents

    def delete_processed_document(self, document_id: str) -> bool:
        """
        Delete processed document

        Args:
            document_id: Document identifier

        Returns:
            True if deleted successfully
        """
        file_path = self.processed_dir / f"{document_id}.json"

        if file_path.exists():
            file_path.unlink()
            print(f"✓ Deleted document: {document_id}")
            return True

        return False


def main():
    """Test the document processor"""
    print("=== Document Processor Test ===\n")

    processor = DocumentProcessor()

    # Check if we have any test documents
    test_files = list(Path("data/uploads").glob("*.*"))

    if not test_files:
        print("No test documents found in data/uploads/")
        print("Add a PDF or DOCX file to test the processor")
        return

    # Process first file
    test_file = test_files[0]
    print(f"Testing with: {test_file.name}\n")

    result = processor.process_document(str(test_file))

    if result["success"]:
        print(f"\n✅ Processing successful!")
        print(f"  Document ID: {result['document_id']}")
        print(f"  Chunks created: {result['num_chunks']}")
        print(f"  Processing time: {result['processing_time']:.2f}s")

        # Save result
        save_path = processor.save_processed_document(result)
        print(f"  Saved to: {save_path}")

        # Show first chunk
        if result["chunks"]:
            first_chunk = result["chunks"][0]
            print(f"\n  First chunk preview:")
            print(f"    {first_chunk['text'][:200]}...")
    else:
        print(f"\n❌ Processing failed: {result['error']}")


if __name__ == "__main__":
    main()